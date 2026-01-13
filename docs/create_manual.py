"""
Script to create the Acumatica Inventory Scanner Manual with embedded screenshots.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Path to images
IMAGES_DIR = os.path.join(os.path.dirname(__file__), 'images')

def add_heading_style(doc):
    """Add custom heading styles"""
    pass  # Use default styles

def add_horizontal_line(doc):
    """Add a horizontal line"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    
def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_step_box(doc, step_num, text):
    """Add a styled step box"""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    cell.text = f"Step {step_num}: {text}"
    set_cell_shading(cell, 'E0F7FA')  # Light cyan
    paragraph = cell.paragraphs[0]
    run = paragraph.runs[0]
    run.bold = True
    run.font.size = Pt(11)
    doc.add_paragraph()  # spacing

def add_image_with_caption(doc, image_path, caption, width=5.5):
    """Add an image with caption"""
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(width))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add caption
        caption_para = doc.add_paragraph()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption_para.add_run(caption)
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(100, 100, 100)
    else:
        p = doc.add_paragraph(f"[Image not found: {image_path}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()  # spacing

def create_manual():
    doc = Document()
    
    # Title
    title = doc.add_heading('Acumatica Inventory Scanner', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('User Manual')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 128, 128)
    
    doc.add_paragraph()
    
    company = doc.add_paragraph()
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = company.add_run('Developed by AcuPower LTD')
    run.font.size = Pt(12)
    run.italic = True
    
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Introduction',
        '2. Prerequisites',
        '3. Acumatica Configuration',
        '   3.1. Navigating to Connected Applications',
        '   3.2. Creating an OAuth Application',
        '   3.3. Adding a Client Secret',
        '   3.4. Saving Your Credentials',
        '4. Using the Mobile App',
        '5. Troubleshooting',
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # Section 1: Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'The Acumatica Inventory Scanner is a modern mobile barcode scanning application '
        'designed for Acumatica ERP inventory management. Built with .NET MAUI, it provides '
        'cross-platform deployment on Android and iOS devices.'
    )
    doc.add_paragraph()
    doc.add_paragraph('Key Features:')
    features = [
        'Real-time Barcode Scanning - Fast camera-based barcode detection',
        'Inventory Lookup - Instantly search and view stock item details',
        'OAuth 2.0 Authentication - Secure API access to Acumatica',
        'Settings Persistence - Save credentials for quick re-login',
        'Modern Dark Theme - Industrial-inspired UI design',
        'Cross-Platform - Works on Android and iOS',
    ]
    for feature in features:
        doc.add_paragraph(f'• {feature}')
    
    doc.add_page_break()
    
    # Section 2: Prerequisites
    doc.add_heading('2. Prerequisites', level=1)
    doc.add_paragraph('Before using this app, ensure you have:')
    prereqs = [
        'Acumatica ERP Instance (version 20.2 or later)',
        'User Account with API access permissions',
        'OAuth Connected Application configured in Acumatica',
    ]
    for i, prereq in enumerate(prereqs, 1):
        doc.add_paragraph(f'{i}. {prereq}')
    
    doc.add_page_break()
    
    # Section 3: Acumatica Configuration
    doc.add_heading('3. Acumatica Configuration', level=1)
    doc.add_paragraph(
        'This section guides you through configuring Acumatica to allow the mobile app '
        'to connect using OAuth 2.0 authentication.'
    )
    
    # 3.1 Navigating to Connected Applications
    doc.add_heading('3.1. Navigating to Connected Applications', level=2)
    
    add_step_box(doc, 1, 'Log in to your Acumatica instance as an administrator.')
    add_image_with_caption(doc, 
        os.path.join(IMAGES_DIR, '01-login-page.png'),
        'Figure 1: Acumatica Login Page - Enter your administrator credentials')
    
    add_step_box(doc, 2, 'Click on "More Items" in the left navigation menu.')
    add_image_with_caption(doc,
        os.path.join(IMAGES_DIR, '02-navigation-menu.png'),
        'Figure 2: Main Navigation Menu - The "More Items" option is at the bottom of the left sidebar')
    
    add_step_box(doc, 3, 'Select "Integration" from the expanded menu.')
    add_image_with_caption(doc,
        os.path.join(IMAGES_DIR, '03-integration-menu.png'),
        'Figure 3: Integration Menu - Shows various integration options')
    
    add_step_box(doc, 4, 'Click on "Connected Applications" under Preferences.')
    add_image_with_caption(doc,
        os.path.join(IMAGES_DIR, '04-integration-full-menu.png'),
        'Figure 4: Full Integration Menu - Connected Applications is under Preferences')
    
    doc.add_page_break()
    
    # 3.2 Creating an OAuth Application
    doc.add_heading('3.2. Creating an OAuth Application', level=2)
    
    add_step_box(doc, 1, 'In the Connected Applications screen, click the "+" button to create a new record.')
    add_image_with_caption(doc,
        os.path.join(IMAGES_DIR, '05-connected-applications.png'),
        'Figure 5: Connected Applications Screen - Click "+" to add a new application')
    
    doc.add_paragraph('Fill in the following fields:')
    
    # Create a table for the fields
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    headers = ['Field', 'Value']
    rows_data = [
        ('Client Name', 'InventoryScanner (or your preferred name)'),
        ('Active', 'Checked ✓'),
        ('Flow', 'Resource Owner Password Credentials'),
        ('Plug-In', 'No Plug-In'),
    ]
    
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, '008080')
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    for i, (field, value) in enumerate(rows_data, 1):
        table.rows[i].cells[0].text = field
        table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    add_image_with_caption(doc,
        os.path.join(IMAGES_DIR, 'step2-create-app.png'),
        'Figure 6: Creating the InventoryScanner OAuth Application')
    
    doc.add_page_break()
    
    # 3.3 Adding a Client Secret
    doc.add_heading('3.3. Adding a Client Secret', level=2)
    
    add_step_box(doc, 1, 'Click on the "SECRETS" tab in the Connected Applications form.')
    add_step_box(doc, 2, 'Click "ADD SHARED SECRET" button.')
    add_step_box(doc, 3, 'Enter a description (e.g., "Mobile App Secret").')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('⚠️ IMPORTANT: ')
    run.bold = True
    run.font.color.rgb = RGBColor(200, 0, 0)
    run = p.add_run('Copy and save the generated secret value immediately! The secret is only shown once and cannot be retrieved later.')
    
    doc.add_paragraph()
    add_image_with_caption(doc,
        os.path.join(IMAGES_DIR, 'step3-add-secret.png'),
        'Figure 7: Adding a Shared Secret - Note the masked value in the Secrets grid')
    
    doc.add_page_break()
    
    # 3.4 Saving Your Credentials
    doc.add_heading('3.4. Saving Your Credentials', level=2)
    
    add_step_box(doc, 1, 'Press Ctrl+S to save the Connected Application.')
    add_step_box(doc, 2, 'Note down the following values for the mobile app:')
    
    doc.add_paragraph()
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    cred_data = [
        ('Credential', 'Example'),
        ('Client ID', 'C6ECE655-8FE3-5C1F-C7C8-3309E724BA61@Company'),
        ('Client Secret', '(The value you copied when creating the secret)'),
    ]
    
    for i, (field, value) in enumerate(cred_data):
        table.rows[i].cells[0].text = field
        table.rows[i].cells[1].text = value
        if i == 0:
            set_cell_shading(table.rows[i].cells[0], '008080')
            set_cell_shading(table.rows[i].cells[1], '008080')
            for cell in table.rows[i].cells:
                run = cell.paragraphs[0].runs[0]
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    doc.add_paragraph()
    add_image_with_caption(doc,
        os.path.join(IMAGES_DIR, 'step4-credentials.png'),
        'Figure 8: Completed OAuth Application with Client ID and Secret configured')
    
    doc.add_page_break()
    
    # Section 4: Using the Mobile App
    doc.add_heading('4. Using the Mobile App', level=1)
    
    doc.add_heading('4.1. First Launch Setup', level=2)
    doc.add_paragraph('When you first open the app, you need to configure the connection settings:')
    
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    
    app_fields = [
        ('Field', 'Description', 'Example'),
        ('Instance URL', 'Your Acumatica site URL', 'https://mycompany.acumatica.com/MySite'),
        ('Username', 'Your Acumatica username', 'admin'),
        ('Password', 'Your Acumatica password', '****'),
        ('Tenant', 'Optional - leave empty for single-tenant', ''),
        ('API Version', 'From the /entity endpoint', '24.200.001'),
        ('Client ID', 'OAuth Client ID from Step 3.4', 'GUID@Company'),
        ('Client Secret', 'OAuth Secret from Step 3.3', 'your-secret-key'),
    ]
    
    for i, row_data in enumerate(app_fields):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = cell_text
            if i == 0:
                set_cell_shading(table.rows[i].cells[j], '008080')
                run = table.rows[i].cells[j].paragraphs[0].runs[0]
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    doc.add_paragraph()
    
    doc.add_heading('4.2. Scanning Barcodes', level=2)
    doc.add_paragraph('To scan inventory items:')
    steps = [
        'Point the camera at a barcode - Position it within the scanning frame',
        'Hold steady - The red scanning line indicates the detection area',
        'Automatic detection - The barcode is recognized and searched automatically',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_heading('4.3. Search Results', level=2)
    doc.add_paragraph('After scanning, the app displays:')
    results = [
        'Item ID - Acumatica Inventory ID',
        'Description - Item description',
        'Availability - Current stock levels',
        'Warehouse Location - Where the item is stored',
    ]
    for result in results:
        doc.add_paragraph(f'• {result}')
    
    doc.add_page_break()
    
    # Section 5: Troubleshooting
    doc.add_heading('5. Troubleshooting', level=1)
    
    issues = [
        ('"401 Unauthorized" Error', [
            'OAuth credentials may be incorrect or expired',
            'Verify Client ID and Secret in Acumatica',
            'Check that the Connected Application is Active',
        ]),
        ('"404 Not Found" Error', [
            'API version mismatch',
            'The endpoint uses StockItem, not InventoryItem',
            'Try a different API version from the /entity endpoint',
        ]),
        ('"Connection Failed"', [
            'Check network connectivity',
            'Verify the instance URL is correct',
            'Ensure no VPN or firewall is blocking access',
        ]),
        ('Scanner Not Detecting', [
            'Ensure camera permissions are granted',
            'Hold device steady with good lighting',
            'Barcode must be within the scanning frame',
        ]),
    ]
    
    for issue_title, solutions in issues:
        p = doc.add_paragraph()
        run = p.add_run(issue_title)
        run.bold = True
        run.font.size = Pt(12)
        
        for solution in solutions:
            doc.add_paragraph(f'• {solution}')
        doc.add_paragraph()
    
    # Footer
    doc.add_page_break()
    doc.add_heading('Support', level=1)
    doc.add_paragraph('Created by AcuPower LTD')
    doc.add_paragraph('Website: https://acupowererp.com')
    doc.add_paragraph('Email: support@acupowererp.com')
    
    # Save the document
    output_path = os.path.join(os.path.dirname(__file__), 'Acumatica_Inventory_Scanner_Manual.docx')
    doc.save(output_path)
    print(f'Manual created successfully: {output_path}')
    return output_path

if __name__ == '__main__':
    create_manual()
